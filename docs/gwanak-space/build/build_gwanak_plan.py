from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/yangjinhwan/Projects/idle-game-generator")
OUT_DIR = ROOT / "docs/gwanak-space"
DOCX_OUT = OUT_DIR / "gwanak-svalley-business-plan.docx"
DIAGRAM_DIR = ROOT / "docs/oracle-miracle/assets/diagrams"

FONT = "Malgun Gothic"
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
LINE = "D9DEE5"


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None, italic: bool | None = None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para_format(p, before=0, after=6, line=1.18, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=6,
             line=1.18, align=None, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    set_para_format(p, before=before, after=after, line=line, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        set_para_format(p, before=14, after=7, line=1.15)
        size = 15
        color = BLUE
    elif level == 2:
        set_para_format(p, before=10, after=5, line=1.15)
        size = 12.5
        color = BLUE
    else:
        set_para_format(p, before=8, after=4, line=1.15)
        size = 11
        color = RGBColor(60, 60, 60)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_para_format(p, before=0, after=4, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, size=9.3, bold=False, color=INK, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.text = ""
    set_para_format(p, after=0, line=1.15, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_row_cant_split(table.rows[0])
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, h, size=9.0, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=font_size)
    add_para(doc, "", after=2)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.7)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, color="B7CBE0", sz="8")
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = cell.paragraphs[0]
    set_para_format(p, after=3, line=1.15)
    r = p.add_run(title)
    set_run_font(r, size=10.4, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    set_para_format(p2, after=0, line=1.18)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    add_para(doc, "", after=2)


def add_captioned_image(doc, image_path, caption, width=6.35):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    set_para_format(p, before=6, after=4, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(caption)
    set_run_font(r, size=9, bold=True, color=MUTED)
    p_img = doc.add_paragraph()
    set_para_format(p_img, after=7, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_img.add_run().add_picture(str(image_path), width=Inches(width))


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 11, RGBColor(60, 60, 60)),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    for name in ("List Bullet", "List Bullet 2"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    # A4 override for Korean public-sector application PDFs.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    set_para_format(header, after=0, line=1.0)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("관악S밸리 창업 공간 입주 사업계획서 | ㈜유니온잼")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(footer, after=0, line=1.0)
    r = footer.add_run("UNIONJAM Inc.")
    set_run_font(r, size=8.5, color=MUTED)

    add_para(doc, "2026년 제2차 관악S밸리 창업 공간 신규 입주기업 모집", size=10.5,
             bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "사업계획서", size=24, bold=True, color=RGBColor(0, 0, 0),
             align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "AI 기반 게임 제작 하네스 플랫폼 IDLE Forge 및 글로벌 모바일 게임 사업화",
             size=12.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    add_table(
        doc,
        ["구분", "내용"],
        [
            ["신청기업", "㈜유니온잼 (UNIONJAM Inc.)"],
            ["창업아이템", "AI 게임 제작 하네스 플랫폼 IDLE Forge + 검증 사례 게임 닌자 서바이벌 2"],
            ["입주 희망", "낙성벤처창업센터 독립형 5인실 우선, 평가 결과에 따른 독립형 4인실 배정 수용"],
            ["입주 목적", "상시 개발·검증·채용·투자 대응이 가능한 AI 네이티브 게임 제작 거점 구축"],
            ["핵심 역량", "자체 Unity/.NET 실시간 게임 엔진, AI 콘텐츠 양산 워크플로우, 글로벌 라이브 게임 운영 경험"],
        ],
        [1.28, 5.38],
        header_fill=LIGHT_BLUE,
        font_size=9.7,
    )

    add_callout(
        doc,
        "요약",
        "㈜유니온잼은 AI 콘텐츠 생성, 자체 게임 엔진, 라이브 운영 파이프라인을 결합한 IDLE Forge를 개발하고 있습니다. "
        "관악S밸리 독립형 공간은 대표·서울대학교 출신 CTO·핵심 개발/아트 인력이 상시 협업하며 제품화, 채용, 투자유치, "
        "서울대 공학 컨설팅 연계를 동시에 추진하기 위한 실행 거점입니다.",
    )

    add_heading(doc, "1. 기업 및 사업 개요", 1)
    add_para(
        doc,
        "㈜유니온잼은 글로벌 모바일 게임 라이브 운영 경험을 바탕으로, 10인 이하 소규모 스튜디오도 대형 게임사 수준의 "
        "콘텐츠 운영 빈도와 서비스 안정성을 확보할 수 있도록 돕는 AI 기반 게임 제작 하네스 플랫폼을 개발하고 있습니다. "
        "핵심 제품인 IDLE Forge는 자연어 기반 콘텐츠 생성, 게임 데이터 검증, Unity/.NET 기반 실시간 서버 엔진, "
        "콘텐츠 핫패치 및 라이브 KPI 운영을 하나의 워크플로우로 연결합니다.",
    )
    add_para(
        doc,
        "당사는 2026년 설립 법인이지만, 운영 이관 및 인수 자산을 통해 글로벌 175개국 서비스 경험을 보유한 닌자 IP 게임 2종을 "
        "운영하고 있으며, IDLE Forge의 성능을 자체 게임에서 먼저 검증한 뒤 외부 인디·중소 게임사로 확장하는 단계적 사업화를 추진합니다.",
    )

    doc.add_page_break()
    add_table(
        doc,
        ["항목", "현재 보유 자산 및 성과"],
        [
            ["기술 자산", "Unity 클라이언트 + .NET 서버 기반 idlez 엔진, 콘텐츠 컴파일·검증·배포 하네스"],
            ["AI 워크플로우", "도메인 특화 생성 스킬 30종 이상, 스키마 검증, 밸런스 리뷰, AI 아트/음향/다국어 제작 흐름"],
            ["운영 자산", "닌자키우기·닌자 서바이벌 등 글로벌 라이브 게임 운영 경험과 유저/매출 데이터 기반 검증 환경"],
            ["사업화 방향", "단기: 자체 게임 글로벌 출시 및 매출화, 중장기: IDLE Forge B2B SaaS·라이선스 공급"],
        ],
        [1.45, 5.2],
        font_size=9.4,
    )

    add_heading(doc, "2. 입주 필요성 및 공간 적합성", 1)
    add_para(
        doc,
        "IDLE Forge는 일반 사무형 사업보다 개발·검증·운영이 동시에 일어나는 기술집약형 사업입니다. AI가 생성한 게임 데이터는 "
        "기획 검토, 밸런스 시뮬레이션, Unity 클라이언트 적용, 서버 검증, 라이브 패치 테스트를 반복해야 하며, 이 과정에는 "
        "상시 회의와 기기 테스트, 빌드 검증, 외부 멘토링·투자 미팅 대응이 함께 필요합니다. 단순 비상주 주소나 임시 공유좌석만으로는 "
        "개발 보안, 업무 집중도, 채용 안정성, 장비 운용 측면에서 한계가 큽니다.",
    )
    add_para(
        doc,
        "관악S밸리 독립형 사무실은 당사가 현재 필요한 '개발 거점'의 성격과 정확히 맞습니다. 대표와 CTO를 포함한 핵심 인력은 "
        "관악구 및 서울 서남권 생활권을 중심으로 활동하고 있어 선정 직후 상주 운영이 가능하며, CTO가 서울대학교 출신 기술 인력이라는 "
        "점은 서울대 공학 컨설팅센터, 관악S밸리 기술 네트워크와의 실질적 접점을 만드는 강점입니다. 이는 단순한 지리적 편의가 아니라 "
        "공간 이용률, 채용 속도, 기술 고도화, 투자 대응력을 높이는 실행 조건입니다.",
    )

    add_table(
        doc,
        ["공간 활용 영역", "구체적 사용 계획"],
        [
            ["상시 개발 좌석", "대표·CTO·개발/기획 인력의 IDLE Forge 플랫폼, Unity 클라이언트, .NET 서버 개발"],
            ["AI 콘텐츠 검수", "생성 콘텐츠의 기획 적합성, IP 유사성, 밸런스, 다국어 문구 검수 및 주간 리뷰"],
            ["테스트·장비 운용", "모바일 기기, 빌드 서버, 운영 대시보드, QA 체크리스트를 고정 배치해 반복 검증"],
            ["채용·멘토링", "신규 개발/기획/데이터 인력 온보딩, 서울대 공학 컨설팅 및 외부 멘토링 대응"],
            ["투자·B2B 대응", "데모데이, 오픈이노베이션, 라이선시 미팅을 위한 제품 시연 및 자료 준비"],
        ],
        [1.45, 5.2],
        font_size=9.35,
    )

    add_heading(doc, "3. 기술·서비스 경쟁력", 1)
    add_heading(doc, "3-1. IDLE Forge 구조", 2)
    add_para(
        doc,
        "IDLE Forge는 콘텐츠 양산, 엔진, 라이브 운영의 세 레이어로 구성됩니다. 콘텐츠 레이어는 자연어 입력을 게임 데이터로 변환하고 "
        "스키마·밸런스·드롭률·번역 문구를 검증합니다. 엔진 레이어는 4년 이상 운영 검증된 Unity/.NET 기반 실시간 게임 구조를 활용합니다. "
        "운영 레이어는 콘텐츠 컴파일, 빌드 검증, 패치 배포, KPI 분석을 반복 가능한 파이프라인으로 묶습니다.",
    )
    add_captioned_image(
        doc,
        DIAGRAM_DIR / "04-idle-forge-architecture-v2.png",
        "[그림 1] IDLE Forge 3-레이어 아키텍처",
        width=6.2,
    )
    add_heading(doc, "3-2. 차별성", 2)
    add_bullet(doc, "단편 AI 도구가 아니라 AI 콘텐츠 양산, 자체 엔진, 라이브 운영을 하나의 제품화 흐름으로 결합합니다.")
    add_bullet(doc, "자체 라이브 게임이 플랫폼의 실시간 검증 환경이므로, 외부 고객 확보 전에도 실제 KPI 기반 개선 사이클을 만들 수 있습니다.")
    add_bullet(doc, "AI가 만든 결과물을 곧바로 게임에 넣는 방식이 아니라 스키마, 밸런스, IP, 다국어, 빌드 검증을 통과시키는 하네스 구조입니다.")
    add_bullet(doc, "닌자 서바이벌 2를 검증 사례 게임으로 개발하여 플랫폼 성능과 콘텐츠 운영 효율을 시장에서 증명합니다.")

    add_heading(doc, "4. 시장성 및 사업화 전략", 1)
    add_para(
        doc,
        "1차 시장은 글로벌 모바일 전략·방치·서바이벌 장르입니다. 이 시장의 핵심 경쟁력은 단발성 아이디어보다 콘텐츠의 양, 운영 속도, "
        "다국어 대응, 이벤트 반복 능력입니다. IDLE Forge는 이 진입 장벽을 낮춰 자체 게임의 출시 속도를 높이고, 향후 외부 인디·중소 게임사가 "
        "동일한 워크플로우를 사용할 수 있도록 공급하는 것을 목표로 합니다.",
    )
    add_table(
        doc,
        ["단계", "기간", "사업화 목표", "핵심 지표"],
        [
            ["1단계", "2026.07~2026.12", "독립형 입주 후 IDLE Forge v1 개발 거점화, 닌자 서바이벌 2 베타 준비", "상주 체계, 주간 빌드, 핵심 파이프라인 안정화"],
            ["2단계", "2027.01~2027.06", "자체 게임 출시·운영 데이터 확보, 투자/퍼블리싱 미팅", "리텐션, ARPDAU, 콘텐츠 생산 속도, 데모데이 참여"],
            ["3단계", "2027 하반기~", "B2B 라이선스 문서화 및 1호 라이선시 발굴", "SDK·온보딩 문서, PoC 계약, 매출분배 모델"],
        ],
        [0.8, 1.22, 3.25, 1.38],
        font_size=8.9,
    )
    add_para(
        doc,
        "수익모델은 자체 게임 매출과 플랫폼 공급 매출의 이중 구조입니다. 단기에는 IDLE Forge로 제작한 게임의 IAP·IAA 매출을 통해 제품성을 "
        "검증하고, 중장기에는 인디·중소 게임사를 대상으로 셋업/온보딩 비용과 매출분배형 라이선스를 결합합니다. 관악S밸리 입주 기간에는 "
        "매출 확대 자체보다 제품 안정화, 투자 유치, 외부 파트너 검증에 집중합니다.",
    )

    add_heading(doc, "5. 조직 역량 및 상주 계획", 1)
    add_para(
        doc,
        "당사의 경쟁력은 기술과 콘텐츠 운영 경험을 동시에 가진 핵심 인력에 있습니다. CTO는 서울대학교 출신으로 게임 개발·서버 백엔드 "
        "10년 이상, idlez 엔진 직접 개발, 글로벌 라이브 게임 운영 및 투자유치 경험을 보유한 핵심 기술 책임자입니다. 대표는 글로벌 게임 "
        "기획·운영, AI 도구 기반 프로토타이핑, 사업화 전략을 총괄합니다. 아트 인력은 닌자 IP 라이브 서비스 전 기간에 참여한 경험을 바탕으로 "
        "AI 아트 워크플로우의 품질 기준을 담당합니다.",
    )
    add_table(
        doc,
        ["역할", "주요 책임", "입주 후 활용"],
        [
            ["대표/CEO", "사업화, 제품 기획, KPI·투자·B2B 전략", "상주 총괄, 관악S밸리 프로그램 참여, 데모데이/투자 대응"],
            ["CTO", "IDLE Forge 아키텍처, idlez 엔진, 서버/클라우드 전환", "상시 개발 리드, 서울대 공학 컨설팅 연계, 기술 검증"],
            ["아트/콘텐츠", "AI 아트 워크플로우, IP 품질관리, 콘텐츠 자산 제작", "생성 결과물 검수, 게임 리소스 생산, UI/UX 산출물 제작"],
            ["채용 예정 인력", "Unity 클라이언트, 콘텐츠 기획, 데이터 운영, B2B 문서화", "입주 공간에서 온보딩 및 주간 개발 스프린트 참여"],
        ],
        [1.25, 2.72, 2.68],
        font_size=9.0,
    )
    add_para(
        doc,
        "입주 후에는 대표와 CTO 중심의 상시 상주 체계를 기본으로 하며, 개발·아트·기획 인력이 주간 스프린트 단위로 사무실에 결합합니다. "
        "호실 규모 대비 근무 인원이 부족해 보이지 않도록, 2026년 하반기 Unity 클라이언트·콘텐츠 기획·데이터 운영 인력을 순차 채용하고 "
        "입주 공간을 온보딩 거점으로 활용할 계획입니다.",
    )

    add_heading(doc, "6. 관악S밸리 프로그램 활용 계획", 1)
    add_table(
        doc,
        ["지원 프로그램", "활용 계획", "기대 산출물"],
        [
            ["서울대 공학 컨설팅", "AI 생성 콘텐츠 검증, 게임 서버 확장성, 데이터 파이프라인 구조에 대한 기술 자문", "기술 리스크 감소, 플랫폼 v1 구조 고도화"],
            ["투자유치 프로그램", "데모데이·IR 멘토링을 통해 Seed/전략투자 라운드 준비", "IR 덱, 제품 데모, 투자자 미팅"],
            ["오픈이노베이션", "게임사·콘텐츠사·클라우드/AI 기업과 PoC 협업 발굴", "B2B 라이선시 후보 및 공동 실증"],
            ["전시회·홍보", "CES/MWC/WIS 등 기술 전시와 게임/콘텐츠 행사 연계 홍보", "관악S밸리 기반 AI 게임 제작 사례 확산"],
            ["네트워킹", "입주기업·서울대·관악권 인재 네트워크 기반 채용 및 협업", "핵심 인력 확보, 지역 생태계 기여"],
        ],
        [1.35, 3.0, 2.3],
        font_size=8.9,
    )
    add_para(
        doc,
        "당사는 입주기업 프로그램 참여율 70% 이상을 내부 운영 목표로 설정하고, 매월 참여 실적과 사업 지표를 함께 관리하겠습니다. "
        "관악S밸리 공간을 단순 임대 공간이 아니라 기술 검증, 네트워크, 투자 유치, 채용의 실행 플랫폼으로 활용하는 것이 목표입니다.",
    )

    add_heading(doc, "7. 입주 후 12개월 추진 일정", 1)
    add_table(
        doc,
        ["기간", "주요 실행", "공간 활용", "성과 지표"],
        [
            ["2026.07", "입주·주소 이전·개발 장비 세팅", "독립형 좌석/테스트 환경 구축", "30일 내 입주 및 비용 납부"],
            ["2026.08~09", "IDLE Forge 핵심 워크플로우 고도화", "주간 개발 스프린트, 콘텐츠 검수 회의", "주 1회 빌드, 생성/검증 로그 확보"],
            ["2026.10~12", "닌자 서바이벌 2 베타 및 데모 준비", "QA, 외부 멘토링, 투자 미팅", "베타 빌드, IR 자료, 사용자 지표"],
            ["2027.01~03", "글로벌 운영 데이터 분석 및 제품 개선", "데이터 리뷰, 채용 온보딩", "리텐션·매출·콘텐츠 생산성 지표"],
            ["2027.04~06", "B2B 공급 문서화 및 후속 투자 준비", "데모데이·라이선시 미팅", "SDK/문서 초안, PoC 후보, 후속 입주 평가 대응"],
        ],
        [1.05, 2.1, 1.9, 1.6],
        font_size=8.65,
    )

    add_heading(doc, "8. 지역 생태계 기여 및 기대효과", 1)
    add_para(
        doc,
        "유니온잼은 관악S밸리 입주를 통해 관악구 기반 AI 게임 제작 스타트업 사례를 만들고자 합니다. 특히 서울대학교와 인접한 관악S밸리의 "
        "기술 네트워크, 청년 인재, 창업 보육 인프라를 활용하여 AI 게임 제작·운영 기술을 고도화하고, 채용과 협업을 지역 생태계 안에서 "
        "확장하겠습니다.",
    )
    add_bullet(doc, "청년 개발자·게임 기획자·데이터 운영 인력 채용을 통해 관악권 기술 일자리 창출에 기여합니다.")
    add_bullet(doc, "AI 기반 콘텐츠 제작·검증 노하우를 입주기업 네트워킹과 데모데이를 통해 공유할 수 있는 사례 기업이 됩니다.")
    add_bullet(doc, "글로벌 게임 출시와 B2B 플랫폼 공급을 통해 관악S밸리의 기술창업 브랜드를 국내외 게임·콘텐츠 시장에 노출합니다.")
    add_bullet(doc, "선정 후 본점 또는 연구소 소재지 이전, 입주 공간 상시 활용, 지원 프로그램 참여 의무를 성실히 이행합니다.")

    doc.add_page_break()
    add_heading(doc, "9. 공고 평가항목 대응 요약", 1)
    add_table(
        doc,
        ["평가항목", "본 계획서 대응"],
        [
            ["조직역량", "서울대학교 출신 CTO, 자체 엔진 개발 경험, 글로벌 라이브 게임 운영 인력, 관악권 상주 가능 체계"],
            ["기술·서비스 경쟁력", "AI 콘텐츠 생성 + 검증 하네스 + Unity/.NET 엔진 + 라이브 운영 파이프라인 수직 통합"],
            ["시장성", "자체 게임 매출화 후 IDLE Forge B2B 공급으로 확장하는 이중 사업모델"],
            ["사업성과 및 성장성", "운영 이관 게임 자산, AI 워크플로우 PoC, 정부지원사업 평가 통과 경험, 채용·투자 계획"],
            ["공간활용도", "독립형 사무실 기반 상시 개발, 테스트 장비 운용, 채용 온보딩, 멘토링·투자·B2B 미팅 수행"],
        ],
        [1.55, 5.1],
        header_fill=LIGHT_BLUE,
        font_size=9.1,
    )
    add_callout(
        doc,
        "입주 이행 약속",
        "선정 시 입주개시일로부터 30일 이내 입주, 본점 또는 연구소 소재지 이전, 사용료·관리비 납부, "
        "입주기업 지원 프로그램 70% 이상 참여를 내부 운영 기준으로 삼고 성실히 이행하겠습니다.",
    )

    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build()
