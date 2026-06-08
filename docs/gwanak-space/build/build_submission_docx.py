from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/yangjinhwan/Projects/idle-game-generator")
SRC = ROOT / "docs/gwanak-space/gwanak-svalley-business-plan-final.md"
OUT = ROOT / "docs/gwanak-space/gwanak-svalley-business-plan-submission.docx"

FONT = "Malgun Gothic"
INK = RGBColor(28, 28, 28)
MUTED = RGBColor(85, 85, 85)
BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F8"
LINE = "D8DEE8"
TABLE_WIDTH_DXA = 9740
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 64, "start": 120, "bottom": 64, "end": 120}


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_para(p, before=0, after=6, line=1.22, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), FONT)
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.append(size)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(r_pr)
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_runs(p, text, size=10.5, color=INK):
    # Support **bold**, `code`, and bare URLs.
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|https?://\S+)")
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_run_font(r, size=size, color=color)
        tok = m.group(0)
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            set_run_font(r, size=size, bold=True, color=color)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            set_run_font(r, size=size, bold=True, color=BLUE)
        elif tok.startswith("http"):
            add_hyperlink(p, tok, tok)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, size=size, color=color)


def cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=CELL_MARGIN["top"], start=CELL_MARGIN["start"], bottom=CELL_MARGIN["bottom"], end=CELL_MARGIN["end"]):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_border(cell, color=LINE, sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width_node(table, width_dxa=TABLE_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA):
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def table_widths(col_count):
    if col_count == 1:
        return [TABLE_WIDTH_DXA]
    if col_count == 2:
        return [2300, TABLE_WIDTH_DXA - 2300]
    if col_count == 3:
        return [1750, 3450, TABLE_WIDTH_DXA - 5200]
    if col_count == 4:
        return [1900, 2400, 2700, TABLE_WIDTH_DXA - 7000]

    base = TABLE_WIDTH_DXA // col_count
    widths = [base] * col_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def set_table_geometry(table, widths):
    table.autofit = False
    set_table_width_node(table)

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_cell_text(cell, text, header=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(cell)
    cell_border(cell)
    p = cell.paragraphs[0]
    p.text = ""
    set_para(p, after=0, line=1.12, align=WD_ALIGN_PARAGRAPH.CENTER if header else None)
    add_inline_runs(p, text, size=8.8 if not header else 8.9, color=INK)
    for r in p.runs:
        set_run_font(r, bold=header)


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    rows = [row + [""] * (col_count - len(row)) for row in rows]
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    widths = table_widths(col_count)
    set_table_geometry(table, widths)
    set_row_cant_split(table.rows[0])
    set_row_repeat_header(table.rows[0])
    for i, val in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell_shading(cell, LIGHT_BLUE if col_count <= 2 else LIGHT_GRAY)
        set_cell_text(cell, val, header=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, val in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], val)
    spacer = doc.add_paragraph()
    set_para(spacer, after=2)


def parse_table(lines, i):
    table_lines = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    data = []
    for line in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        inner = line.strip("|")
        data.append([c.strip() for c in inner.split("|")])
    return data, i


def configure_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    # Korean public application documents are normally A4.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(header, after=0, line=1.0)
    r = header.add_run("관악S밸리 창업 공간 입주 사업계획서 | ㈜유니온잼")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, after=0, line=1.0)
    r = footer.add_run("UNIONJAM Inc.")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    set_para(p, before=20, after=8, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("2026년 제2차 관악S밸리 창업 공간 신규 입주기업 모집")
    set_run_font(r, size=13, bold=True, color=MUTED)

    p = doc.add_paragraph()
    set_para(p, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("사업계획서")
    set_run_font(r, size=26, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    set_para(p, after=12, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AI 기반 게임 제작·운영 플랫폼 IDLE Forge 및 닌자 서바이벌 2")
    set_run_font(r, size=12.6, color=MUTED)

    p = doc.add_paragraph()
    set_para(p, after=12, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("제출기업: 주식회사 유니온잼 | 희망 입주 형태: 독립형 사무실")
    set_run_font(r, size=10.5, bold=True, color=BLUE)


def build():
    md = SRC.read_text(encoding="utf-8")
    lines = md.splitlines()
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    i = 1  # skip H1 in markdown; cover renders title.
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            set_para(p, before=12, after=6, line=1.1)
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, line[3:], size=16, color=BLUE)
            for r in p.runs:
                set_run_font(r, bold=True)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            set_para(p, before=8, after=4, line=1.1)
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, line[4:], size=13, color=BLUE)
            for r in p.runs:
                set_run_font(r, bold=True)
            i += 1
            continue

        if line.startswith("|"):
            table, i = parse_table(lines, i)
            add_table(doc, table)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_para(p, before=0, after=3, line=1.16)
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.194)
            add_inline_runs(p, line[2:], size=10.2)
            i += 1
            continue

        p = doc.add_paragraph()
        set_para(p, before=0, after=4, line=1.18)
        add_inline_runs(p, line, size=10.5)
        i += 1

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
